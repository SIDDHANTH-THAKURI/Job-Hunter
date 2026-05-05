"""
Run once (or after any design change) to create templates/resume.docx and templates/cover.docx.
After running, open either file in Word to adjust styling.
Do NOT delete the {{ }} or {%p %} tags — they are filled at generation time.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = Path(__file__).parent / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)

FONT   = "Calibri"
ACCENT = RGBColor(0x1A, 0x56, 0xDB)
MUTED  = RGBColor(0x55, 0x65, 0x7A)

# A4 content width in twips: 210mm − 2×0.75" margins
# 0.75" = 19.05mm → content = 171.9mm × 56.693 twips/mm ≈ 9745
RT = 9745


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _sp(para, before=0, after=2):
    pPr = para._p.get_or_add_pPr()
    s = OxmlElement("w:spacing")
    s.set(qn("w:before"), str(before * 20))
    s.set(qn("w:after"),  str(after  * 20))
    pPr.append(s)


def _run(para, text, bold=False, size=9.5, italic=False, color=None):
    r = para.add_run(text)
    r.font.name  = FONT
    r.font.size  = Pt(size)
    r.bold       = bold
    r.italic     = italic
    if color:
        r.font.color.rgb = color
    return r


def _loop(doc, tag):
    """
    Jinja2 loop/conditional marker.
    Rendered white at 1pt — invisible in the template.
    docxtpl removes these paragraphs from the rendered output.
    """
    p = doc.add_paragraph()
    _sp(p, before=0, after=0)
    r = p.add_run(tag)
    r.font.name  = FONT
    r.font.size  = Pt(1)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _section(doc, title):
    p = doc.add_paragraph()
    _sp(p, before=8, after=2)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "4")
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "1A56DB")
    bdr.append(btm)
    pPr.append(bdr)
    _run(p, title.upper(), bold=True, size=10.5, color=ACCENT)


def _row(doc, left_expr, right_expr, bold=True, size=9.5):
    """
    Title row: left-aligned entry title + right-aligned date via tab stop.
    Keep left_expr short (title only, not company) to avoid overflow.
    """
    p = doc.add_paragraph()
    _sp(p, before=5, after=0)
    pPr = p._p.get_or_add_pPr()
    tabs_el = OxmlElement("w:tabs")
    t = OxmlElement("w:tab")
    t.set(qn("w:val"), "right")
    t.set(qn("w:pos"), str(RT))
    tabs_el.append(t)
    pPr.append(tabs_el)
    _run(p, left_expr, bold=bold, size=size)
    _run(p, "\t" + right_expr, size=size - 0.5, color=MUTED)


def _sub(doc, expr, size=8.5):
    """Subtitle row — company name or institution, no tab stop, muted."""
    p = doc.add_paragraph()
    _sp(p, before=0, after=1)
    _run(p, expr, size=size, color=MUTED)


def _body(doc, expr, size=9.5, italic=False, muted=False):
    p = doc.add_paragraph()
    _sp(p, after=2)
    _run(p, expr, size=size, italic=italic, color=MUTED if muted else None)


def _bullet(doc, expr, size=9.5):
    p = doc.add_paragraph()
    _sp(p, after=1)
    p.paragraph_format.left_indent       = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.17)
    _run(p, "•  " + expr, size=size)


def _skill_richtext(doc):
    """
    Skills use {{ skill }} — skill is a RichText object passed directly in context.
    Avoids multi-run template issues and the {{ r x.y }} dotted-path Jinja2 bug.
    """
    p = doc.add_paragraph()
    _sp(p, after=2)
    r = p.add_run("{{ skill }}")
    r.font.name = FONT
    r.font.size = Pt(9.5)


# ─── Resume template ──────────────────────────────────────────────────────────

def create_resume_template():
    doc = Document()

    for sec in doc.sections:
        sec.page_width    = Mm(210)
        sec.page_height   = Mm(297)
        sec.top_margin    = Inches(0.65)
        sec.bottom_margin = Inches(0.65)
        sec.left_margin   = Inches(0.75)
        sec.right_margin  = Inches(0.75)

    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(9.5)

    # ── Name block ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, after=0)
    _run(p, "{{ name }}", bold=True, size=16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, after=3)
    _run(p, "{{ title_subtitle }}", size=11, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, after=1)
    _run(p, "{{ contact1 }}", size=8.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, after=5)
    _run(p, "{{ contact2 }}", size=8.5, color=MUTED)

    # ── Career Objective ────────────────────────────────────────────────────
    _section(doc, "Career Objective")
    _body(doc, "{{ career_objective }}")

    # ── Technical Skills — RichText approach avoids formatting bugs ──────────
    _section(doc, "Technical Skills")
    _loop(doc, "{%p for skill in skills %}")
    _skill_richtext(doc)
    _loop(doc, "{%p endfor %}")

    # ── Relevant Experience ─────────────────────────────────────────────────
    _section(doc, "Relevant Experience")
    _loop(doc, "{%p for exp in relevant_experience %}")
    _row(doc,  "{{ exp.title }}",   "{{ exp.period }}")   # title only — stays short
    _sub(doc,  "{{ exp.company }}")                        # company on its own line
    _body(doc, "Role Description: {{ exp.role_description }}", size=9, italic=True, muted=True)
    _loop(doc, "{%p for b in exp.bullets %}")
    _bullet(doc, "{{ b }}")
    _loop(doc, "{%p endfor %}")
    _loop(doc, "{%p endfor %}")

    # ── Additional Experience ───────────────────────────────────────────────
    _loop(doc, "{%p if other_experience %}")
    _section(doc, "Additional Experience")
    _loop(doc, "{%p endif %}")
    _loop(doc, "{%p for exp in other_experience %}")
    _row(doc,  "{{ exp.title }}",   "{{ exp.period }}")
    _sub(doc,  "{{ exp.company }}")
    _body(doc, "Role Description: {{ exp.role_description }}", size=9, italic=True, muted=True)
    _loop(doc, "{%p for b in exp.bullets %}")
    _bullet(doc, "{{ b }}")
    _loop(doc, "{%p endfor %}")
    _loop(doc, "{%p endfor %}")

    # ── Education ───────────────────────────────────────────────────────────
    _section(doc, "Education")
    _loop(doc, "{%p for edu in education %}")
    _row(doc,  "{{ edu.degree }}",       "{{ edu.period }}")
    _sub(doc,  "{{ edu.institution }}")
    _loop(doc, "{%p for b in edu.bullets %}")
    _bullet(doc, "{{ b }}")
    _loop(doc, "{%p endfor %}")
    _loop(doc, "{%p endfor %}")

    # ── Academic Projects ───────────────────────────────────────────────────
    _loop(doc, "{%p if academic_projects %}")
    _section(doc, "Academic Projects")
    _loop(doc, "{%p endif %}")
    _loop(doc, "{%p for proj in academic_projects %}")
    _row(doc,  "{{ proj.name }}",    "{{ proj.period }}")
    _sub(doc,  "{{ proj.context }}")
    _loop(doc, "{%p for b in proj.bullets %}")
    _bullet(doc, "{{ b }}")
    _loop(doc, "{%p endfor %}")
    _loop(doc, "{%p endfor %}")

    # ── Personal Projects ───────────────────────────────────────────────────
    _loop(doc, "{%p if personal_projects %}")
    _section(doc, "Personal Projects")
    _loop(doc, "{%p endif %}")
    _loop(doc, "{%p for proj in personal_projects %}")
    _row(doc,  "{{ proj.name }}",   "{{ proj.period }}")
    _loop(doc, "{%p for b in proj.bullets %}")
    _bullet(doc, "{{ b }}")
    _loop(doc, "{%p endfor %}")
    _loop(doc, "{%p endfor %}")

    # ── Additional Qualifications (short courses, professional development) ───
    _loop(doc, "{%p if additional_education %}")
    _section(doc, "Additional Qualifications")
    _loop(doc, "{%p endif %}")
    _loop(doc, "{%p for edu in additional_education %}")
    _row(doc,  "{{ edu.degree }}",    "{{ edu.period }}")
    _sub(doc,  "{{ edu.institution }}")
    _loop(doc, "{%p for b in edu.bullets %}")
    _bullet(doc, "{{ b }}")
    _loop(doc, "{%p endfor %}")
    _loop(doc, "{%p endfor %}")

    # ── Affiliations ────────────────────────────────────────────────────────
    _loop(doc, "{%p if affiliations %}")
    _section(doc, "Professional Affiliations")
    _loop(doc, "{%p endif %}")
    _loop(doc, "{%p for aff in affiliations %}")
    _bullet(doc, "{{ aff }}")
    _loop(doc, "{%p endfor %}")

    out = TEMPLATES_DIR / "resume.docx"
    doc.save(str(out))
    print(f"  Created: {out}")


# ─── Cover letter template ────────────────────────────────────────────────────

def create_cover_template():
    doc = Document()

    for sec in doc.sections:
        sec.page_width    = Mm(210)
        sec.page_height   = Mm(297)
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(10)

    def _sp_c(para, before=0, after=0):
        pPr = para._p.get_or_add_pPr()
        s = OxmlElement("w:spacing")
        s.set(qn("w:before"), str(before * 20))
        s.set(qn("w:after"),  str(after  * 20))
        pPr.append(s)

    def _line(text, size=10, bold=False, color=None, after=0):
        p = doc.add_paragraph()
        _sp_c(p, after=after)
        _run(p, text, bold=bold, size=size, color=color)

    _line("{{ name }}",     size=12, bold=True)
    _line("{{ contact1 }}", size=9,  color=MUTED)
    _line("{{ contact2 }}", size=9,  color=ACCENT, after=14)
    _line("{{ date }}",     size=10, after=8)
    _line("Hiring Manager", size=10)
    _line("{{ company }}",  size=10, after=14)

    p = doc.add_paragraph()
    _sp_c(p, after=14)
    _run(p, "Re: {{ subject }}", bold=True, size=10)

    _loop(doc, "{%p for para in paragraphs %}")
    p = doc.add_paragraph()
    _sp_c(p, after=8)
    _run(p, "{{ para }}", size=10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _loop(doc, "{%p endfor %}")

    p = doc.add_paragraph()
    _sp_c(p, before=8, after=18)
    _run(p, "Yours sincerely,", size=10)

    p = doc.add_paragraph()
    _sp_c(p)
    _run(p, "{{ name }}", bold=True, size=10)

    out = TEMPLATES_DIR / "cover.docx"
    doc.save(str(out))
    print(f"  Created: {out}")


if __name__ == "__main__":
    print("Creating templates…")
    create_resume_template()
    create_cover_template()
    print("\nDone. Open templates/resume.docx or cover.docx in Word to adjust styling.")
    print("Do NOT delete the {{ }} or {%p %} tags.")
