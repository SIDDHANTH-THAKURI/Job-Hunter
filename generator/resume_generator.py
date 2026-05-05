import json
import os
import re
import sqlite3
from datetime import date
from pathlib import Path

import anthropic
from dotenv import load_dotenv

load_dotenv()

BASE_DIR = Path(__file__).parent.parent
DB_PATH = BASE_DIR / "data" / "jobs.db"
ME_PATH = BASE_DIR / "me.json"
OUTPUT_DIR = BASE_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


def _db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def _profile():
    with open(ME_PATH, encoding="utf-8-sig") as f:
        return json.load(f)


def _clean(s: str) -> str:
    """Replace Unicode dashes and non-breaking spaces with plain ASCII equivalents."""
    return (s
            .replace('–', '-')   # en-dash
            .replace('—', '-')   # em-dash
            .replace('−', '-')   # minus sign
            .replace('‐', '-')   # hyphen
            .replace('‑', '-')   # non-breaking hyphen
            .replace(' ', ' ')   # non-breaking space
            )





_MONTH_ABBR = [
    ("January","Jan"),("February","Feb"),("March","Mar"),("April","Apr"),
    ("May","May"),("June","Jun"),("July","Jul"),("August","Aug"),
    ("September","Sep"),("October","Oct"),("November","Nov"),("December","Dec"),
]


def _fmt_period(s: str) -> str:
    s = _clean(s)
    for full, short in _MONTH_ABBR:
        s = s.replace(full, short)
    return s


def _link(url: str) -> str:
    if url and not url.startswith("http"):
        return "https://" + url
    return url or ""


_URL_RE = re.compile(r'(https?://\S+)')


def _rt_with_links(text: str, tpl, size: int = 19) -> object:
    from docxtpl import RichText
    if not _URL_RE.search(text):
        return text
    rt = RichText()
    for part in _URL_RE.split(text):
        if _URL_RE.fullmatch(part):
            url = part.rstrip(".,;)\"']")
            suffix = part[len(url):]
            rt.add(url, url_id=tpl.build_url_id(url), font="Calibri", size=size,
                   color="1A56DB", underline=True)
            if suffix:
                rt.add(suffix, font="Calibri", size=size)
        elif part:
            rt.add(part, font="Calibri", size=size)
    return rt


# ─── Claude prompt (built dynamically from me.json) ──────────────────────────

def _build_prompt(profile: dict, job: dict) -> str:
    p = profile
    personal = p["personal"]
    portfolio = personal.get("portfolio", "")

    # ── Bullet-count constraints from work_experience ──
    bullet_constraints = []
    for exp in p["work_experience"]:
        if exp.get("show_on_resume"):
            label = exp.get("resume_label") or exp["title"]
            count = exp.get("bullet_count", 4)
            bullet_constraints.append(f"- {label}: exactly {count} bullets, max 22 words each")

    # ── Experience scaffold (split by category) ──
    def _exp_entry(exp):
        title_str = exp["title"]  # clean title only — label is framing context, not for display
        loc = exp.get("location", "").split("(")[0].strip().rstrip(",")
        company_str = f"{exp['company']}, {loc}" if loc else exp["company"]
        period = _fmt_period(exp["period"])
        count = exp.get("bullet_count", 4)
        bullets_ph = ', '.join(['"action verb + detail (max 22 words)"'] * count)
        role_desc = exp.get("framing", "")[:120]
        return (
            f'      {{\n'
            f'        "title": "{title_str}",\n'
            f'        "company": "{company_str}",\n'
            f'        "period": "{period}",\n'
            f'        "role_description": "{role_desc}",\n'
            f'        "bullets": [{bullets_ph}]\n'
            f'      }}'
        )

    rel_entries = [_exp_entry(e) for e in p["work_experience"]
                   if e.get("show_on_resume") and e.get("experience_category", "relevant") == "relevant"]
    oth_entries = [_exp_entry(e) for e in p["work_experience"]
                   if e.get("show_on_resume") and e.get("experience_category") == "other"]
    rel_block = ",\n".join(rel_entries)
    oth_block = ",\n".join(oth_entries)

    # ── Education scaffold (bullets are fixed — not tailored) ──
    def _edu_scaffold(edu_list):
        entries = []
        for edu in edu_list:
            bullets = edu.get("resume_bullets", [edu.get("relevance", "")])
            bullets_json = json.dumps(bullets)
            period = _fmt_period(edu["period"])
            entries.append(
                f'      {{\n'
                f'        "degree": "{edu["degree"]}",\n'
                f'        "institution": "{edu["institution"]}",\n'
                f'        "period": "{period}",\n'
                f'        "bullets": {bullets_json}\n'
                f'      }}'
            )
        return ",\n".join(entries)

    edu_main   = [e for e in p["education"] if e.get("education_category", "main") == "main"]
    edu_extra  = [e for e in p["education"] if e.get("education_category") == "additional"]
    edu_block  = _edu_scaffold(edu_main)
    edu_extra_block = _edu_scaffold(edu_extra)

    # ── Academic projects scaffold ──
    academic_proj_entries = []
    for proj in p["projects"]:
        if proj.get("type", "").startswith("Academic") and proj.get("include_by_default", False):
            name = proj.get("display_name") or proj["name"]
            ctx = proj.get("resume_context", "")
            fixed = proj.get("fixed_bullets", [])
            tailor_count = max(0, 3 - len(fixed))
            tailor = ['"Tailor bullet to job requirements — use details from profile"'] * tailor_count
            all_bullets = [f'"{b}"' for b in fixed]
            all_bullets = tailor + all_bullets  # tailored first, fixed last
            proj_period = _fmt_period(proj["period"])
            academic_proj_entries.append(
                f'      {{\n'
                f'        "name": "{name}",\n'
                f'        "context": "{ctx}",\n'
                f'        "period": "{proj_period}",\n'
                f'        "bullets": [{", ".join(all_bullets)}]\n'
                f'      }}'
            )
    academic_proj_block = ",\n".join(academic_proj_entries)

    # ── Personal projects scaffold ──
    personal_proj_entries = []
    for proj in p["projects"]:
        if not proj.get("type", "").startswith("Academic") and proj.get("include_by_default", False):
            name = proj.get("display_name") or proj["name"]
            fixed = proj.get("fixed_bullets", [])
            tailor_count = max(0, 3 - len(fixed))
            tailor = ['"Tailor bullet to job requirements — use details from profile"'] * tailor_count
            all_bullets = tailor + [f'"{b}"' for b in fixed]
            proj_period = _fmt_period(proj["period"])
            personal_proj_entries.append(
                f'      {{\n'
                f'        "name": "{name}",\n'
                f'        "period": "{proj_period}",\n'
                f'        "bullets": [{", ".join(all_bullets)}]\n'
                f'      }}'
            )
    personal_proj_block = ",\n".join(personal_proj_entries)

    # ── Affiliations ──
    affiliations = []
    for aff in p.get("professional_affiliations", []):
        line = f"{aff['name']} – {aff['type']} | Member ID: {aff['member_id']}"
        affiliations.append(f'      "{line}"')
    affiliations_block = ",\n".join(affiliations)

    # ── Special rules from notes ──
    special_rules = []
    for exp in p["work_experience"]:
        note = exp.get("note", "")
        if "unpaid" in note.lower():
            special_rules.append(f"- Do NOT mention the {exp['company']} internship was unpaid")
    for exp in p["work_experience"]:
        if "catering" in exp.get("title", "").lower() or "kitchen" in exp.get("title", "").lower():
            special_rules.append("- Catering/hospitality roles: frame professionally, highlight transferable skills (communication, process-following, composure under pressure)")
    special_rules_block = "\n".join(special_rules)

    # ── Key project references for cover letter ──
    highlight_projects = [
        proj.get("display_name") or proj["name"]
        for proj in p["projects"]
        if proj.get("include_by_default") and proj.get("type", "").startswith("Academic")
    ]
    proj_ref = highlight_projects[0] if highlight_projects else "a personal project"

    description = (job.get("description") or "")[:3000]
    title = job.get("title") or "the role"
    company = job.get("company") or "the company"

    return f"""You are a professional resume writer. Generate a tailored resume and cover letter.

CANDIDATE PROFILE:
{json.dumps(profile, indent=2)}

TARGET JOB:
Title: {title}
Company: {company}
Location: {job.get("location") or "Sydney, NSW"}
Salary: {job.get("salary") or "Not specified"}
Description:
{description}

Use the FIXED TEMPLATE STRUCTURE below. Every section and every work experience entry is MANDATORY — never remove any.
Tailor the content within each section to align with the job description.

HARD CONSTRAINTS (2-page resume limit):
- career_objective: 80-100 words exactly
{chr(10).join(bullet_constraints)}
- Each project: exactly 3 bullets, max 22 words each
- role_description: max 35 words each
- Skills per category: 8-10 items, most job-relevant first

Return ONLY valid JSON (absolutely no markdown fences, no text outside the JSON):

{{
  "resume": {{
    "title_subtitle": "ROLE TITLE, SPECIALISATION (match the job title, e.g. JUNIOR DEVELOPER, ARTIFICIAL INTELLIGENCE)",
    "career_objective": "3-4 sentences. Mention Accenture + years, Masters degree + grade, top 2-3 skills most relevant to this job. Final sentence: genuine motivation for this specific company domain or mission.",
    "skills": {{
      "Languages": "most relevant first, comma-separated",
      "Frameworks & Libraries": "most relevant first, comma-separated",
      "AI & ML": "most relevant first, comma-separated",
      "Databases": "most relevant first, comma-separated",
      "Cloud & DevOps": "most relevant first, comma-separated",
      "Tools & Practices": "most relevant first, comma-separated"
    }},
    "relevant_experience": [
{rel_block}
    ],
    "other_experience": [
{oth_block}
    ],
    "education": [
{edu_block}
    ],
    "additional_education": [
{edu_extra_block}
    ],
    "academic_projects": [
{academic_proj_block}
    ],
    "personal_projects": [
{personal_proj_block}
    ],
    "affiliations": [
{affiliations_block}
    ]
  }},
  "cover_letter": {{
    "subject": "Application for {title} at {company}",
    "paragraphs": [
      "Opening (3-4 sentences): Genuine enthusiasm. Name the role and company. One-line value proposition referencing Accenture experience and Masters degree.",
      "Body 1 (3-4 sentences): 2-3 specific, evidenced skills or achievements from the profile that directly match the job requirements. Be concrete.",
      "Body 2 (2-3 sentences): Reference {proj_ref} or a personal project as proof of real-world delivery. Connect it to the target role.",
      "Closing (2 sentences): Reiterate strong interest. Call to action and thank you."
    ]
  }}
}}

Rules:
- Australian English spelling throughout
- All bullets start with a strong action verb
- No fabrication — only use facts from the provided candidate profile
- ATS-safe: no tables, columns, text boxes, or special formatting — plain text only
- NEVER use em dashes (—) or en dashes (–) anywhere — use commas, colons, or hyphens (-) instead
- NEVER use smart/curly quotes — use straight quotes only
- Keep all text within standard ASCII/Latin characters where possible
{special_rules_block}
- title_subtitle should match the spirit of the job title, not copy it verbatim"""


def _call_claude(profile: dict, job: dict) -> dict:
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    prompt = _build_prompt(profile, job)
    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )
    try:
        from utils import log_api_call
        log_api_call("claude-sonnet-4-6", "generate", msg.usage.input_tokens, msg.usage.output_tokens)
    except Exception:
        pass

    text = msg.content[0].text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return json.loads(text)


# ─── Template rendering (docxtpl + docx2pdf) ─────────────────────────────────

TEMPLATES_DIR = BASE_DIR / "templates"


def _ensure_templates():
    if not (TEMPLATES_DIR / "resume.docx").exists() or not (TEMPLATES_DIR / "cover.docx").exists():
        import create_templates as ct
        print("[generator] Templates not found — creating them now…")
        ct.create_resume_template()
        ct.create_cover_template()




def _build_context_resume(resume_data: dict, tpl=None) -> dict:
    from docxtpl import RichText

    full_profile = _profile()
    profile = full_profile["personal"]
    portfolio = profile.get("portfolio", "")
    leetcode = profile.get("leetcode", "")

    def _skill_rt(label, value):
        rt = RichText()
        rt.add(f"{label}:  ", bold=True, font="Calibri", size=19)
        rt.add(value, font="Calibri", size=19)
        return rt

    edu_main = [e for e in full_profile.get("education", [])
                if e.get("education_category", "main") == "main"]
    edu_extra = [e for e in full_profile.get("education", [])
                 if e.get("education_category") == "additional"]

    def _edu_entry(e):
        return {
            "degree":      e["degree"],
            "institution": e.get("institution", ""),
            "period":      _fmt_period(e.get("period", "")),
            "bullets":     e.get("resume_bullets", []),
        }

    def _bullets(lst):
        if tpl is None:
            return [_clean(b) for b in lst]
        return [_rt_with_links(_clean(b), tpl) for b in lst]

    # contact2: label + hyperlinked URL for each social profile
    contact2_links = [
        ("LinkedIn",  _link(profile["linkedin"])),
        ("GitHub",    _link(profile["github"])),
        ("Portfolio", _link(portfolio)),
    ]
    if leetcode:
        contact2_links.append(("LeetCode", _link(leetcode)))

    if tpl is not None:
        ct2 = RichText()
        for i, (label, url) in enumerate(contact2_links):
            if i > 0:
                ct2.add("  |  ", font="Calibri", size=17, color="55657A")
            ct2.add(f"{label}: ", font="Calibri", size=17, color="55657A")
            ct2.add(url, url_id=tpl.build_url_id(url), font="Calibri", size=17,
                    color="1A56DB", underline=True)
        contact2 = ct2
    else:
        contact2 = "  |  ".join(f"{l}: {u}" for l, u in contact2_links)

    return {
        "name":             profile["name"].upper(),
        "title_subtitle":   resume_data.get("title_subtitle", ""),
        "contact1":         f"Mobile: {profile['phone']}  |  Email: {profile['email']}  |  Address: {profile['location']}",
        "contact2":         contact2,
        "career_objective": resume_data.get("career_objective", ""),
        "skills": [_skill_rt(k, v) for k, v in resume_data.get("skills", {}).items()],
        "relevant_experience": [
            {
                "title":            e["title"],
                "company":          e.get("company", ""),
                "period":           _fmt_period(e.get("period", "")),
                "role_description": e.get("role_description", ""),
                "bullets":          _bullets(e.get("bullets", [])),
            }
            for e in resume_data.get("relevant_experience", [])
        ],
        "other_experience": [
            {
                "title":            e["title"],
                "company":          e.get("company", ""),
                "period":           _fmt_period(e.get("period", "")),
                "role_description": e.get("role_description", ""),
                "bullets":          _bullets(e.get("bullets", [])),
            }
            for e in resume_data.get("other_experience", [])
        ],
        "education":            [_edu_entry(e) for e in edu_main],
        "additional_education": [_edu_entry(e) for e in edu_extra],
        "academic_projects": [
            {
                "name":    p["name"],
                "context": p.get("context", ""),
                "period":  _fmt_period(p.get("period", "")),
                "bullets": _bullets(p.get("bullets", [])),
            }
            for p in resume_data.get("academic_projects", [])
        ],
        "personal_projects": [
            {
                "name":    p["name"],
                "period":  _fmt_period(p.get("period", "")),
                "bullets": _bullets(p.get("bullets", [])),
            }
            for p in resume_data.get("personal_projects", [])
        ],
        "affiliations": resume_data.get("affiliations", []),
    }


def _fill_resume(resume_data: dict, path: Path):
    from docxtpl import DocxTemplate
    _ensure_templates()
    tpl = DocxTemplate(str(TEMPLATES_DIR / "resume.docx"))
    tpl.render(_build_context_resume(resume_data, tpl))
    tpl.save(str(path))


def _fill_cover(cover_data: dict, job: dict, path: Path):
    from docxtpl import DocxTemplate, RichText
    _ensure_templates()
    profile = _profile()["personal"]
    portfolio = _link(profile.get("portfolio", ""))
    linkedin  = _link(profile["linkedin"])
    tpl = DocxTemplate(str(TEMPLATES_DIR / "cover.docx"))
    ct2 = RichText()
    ct2.add(linkedin, url_id=tpl.build_url_id(linkedin), font="Calibri", size=18,
            color="1A56DB", underline=True)
    ct2.add("  |  ", font="Calibri", size=18, color="55657A")
    ct2.add(portfolio, url_id=tpl.build_url_id(portfolio), font="Calibri", size=18,
            color="1A56DB", underline=True)
    tpl.render({
        "name":      profile["name"],
        "contact1":  f"{profile['phone']}  |  {profile['email']}  |  {profile['location']}",
        "contact2":  ct2,
        "date":      date.today().strftime("%d %B %Y"),
        "company":   job.get("company") or "",
        "subject":   cover_data.get("subject", ""),
        "paragraphs": cover_data.get("paragraphs", []),
    })
    tpl.save(str(path))



def _to_pdf(docx_path: Path, pdf_path: Path):
    from docx2pdf import convert
    convert(str(docx_path), str(pdf_path))


def _make_word_resume(data: dict, path: Path):  # UNUSED — kept so old cached calls don't 500
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.65)
        sec.bottom_margin = Inches(0.65)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    # Content width in twips (8.5" - 0.75" - 0.75" = 7" = 10080 twips)
    RT = 10080

    profile = _profile()["personal"]
    portfolio = profile.get("portfolio", "")

    def _sp(para, before=0, after=2):
        pPr = para._p.get_or_add_pPr()
        s = OxmlElement("w:spacing")
        s.set(qn("w:before"), str(before * 20))
        s.set(qn("w:after"), str(after * 20))
        pPr.append(s)

    def _section(title):
        p = doc.add_paragraph()
        _sp(p, before=8, after=2)
        pPr = p._p.get_or_add_pPr()
        bdr = OxmlElement("w:pBdr")
        btm = OxmlElement("w:bottom")
        btm.set(qn("w:val"), "single")
        btm.set(qn("w:sz"), "4")
        btm.set(qn("w:space"), "1")
        btm.set(qn("w:color"), "1A56DB")
        bdr.append(btm)
        pPr.append(bdr)
        r = p.add_run(title.upper())
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    def _row(left, right, bold=True, size=9.5):
        """Single paragraph: left-aligned title + right-aligned date via tab stop."""
        p = doc.add_paragraph()
        _sp(p, before=3, after=1)
        pPr = p._p.get_or_add_pPr()
        tabs_el = OxmlElement("w:tabs")
        t = OxmlElement("w:tab")
        t.set(qn("w:val"), "right")
        t.set(qn("w:pos"), str(RT))
        tabs_el.append(t)
        pPr.append(tabs_el)
        r1 = p.add_run(left)
        r1.bold = bold
        r1.font.size = Pt(size)
        r2 = p.add_run(f"\t{right}")
        r2.font.size = Pt(size - 0.5)
        r2.font.color.rgb = RGBColor(0x55, 0x65, 0x7A)

    def _body(text, size=9.5, italic=False, muted=False):
        p = doc.add_paragraph()
        _sp(p, after=2)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.italic = italic
        if muted:
            r.font.color.rgb = RGBColor(0x55, 0x65, 0x7A)

    def _bullet(text, size=9.5):
        p = doc.add_paragraph()
        _sp(p, after=1)
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.17)
        r = p.add_run(f"•  {text}")
        r.font.size = Pt(size)

    def _skill_row(label, value, size=9.5):
        p = doc.add_paragraph()
        _sp(p, after=2)
        r1 = p.add_run(f"{label}:    ")
        r1.bold = True
        r1.font.size = Pt(size)
        r2 = p.add_run(value)
        r2.font.size = Pt(size)

    # ── Name block ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, after=0)
    r = p.add_run(profile["name"].upper())
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, after=3)
    r = p.add_run(data.get("title_subtitle", ""))
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x33, 0x44, 0x5A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, after=1)
    r = p.add_run(
        f"Mobile: {profile['phone']}  |  Email: {profile['email']}  |  Address: {profile['location']}"
    )
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x44, 0x55, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, after=5)
    r = p.add_run(
        f"LinkedIn: {profile['linkedin']}  |  GitHub: {profile['github']}  |  Portfolio: {portfolio}"
    )
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x44, 0x55, 0x66)

    # ── Career Objective ────────────────────────────────────────────────
    _section("Career Objective")
    _body(data["career_objective"], size=9.5)

    # ── Technical Skills (no table — plain labelled rows, ATS safe) ─────
    _section("Technical Skills")
    for cat, val in data["skills"].items():
        _skill_row(cat, val)

    # ── Work Experience ─────────────────────────────────────────────────
    _section("Work Experience")
    for exp in data["experience"]:
        _row(f"{exp['title']}  |  {exp['company']}", exp["period"])
        _body(f"Role Description: {exp['role_description']}", size=9, italic=True, muted=True)
        for b in exp.get("bullets", []):
            _bullet(b)

    # ── Education ───────────────────────────────────────────────────────
    _section("Education")
    for edu in data["education"]:
        _row(f"{edu['degree']}  |  {edu['institution']}", edu["period"])
        for b in edu.get("bullets", []):
            _bullet(b)

    # ── Academic Projects ───────────────────────────────────────────────
    if data.get("academic_projects"):
        _section("Academic Projects")
        for proj in data["academic_projects"]:
            _row(
                f"{proj['name']}  |  {proj.get('context', '')}",
                proj.get("period", ""),
            )
            for b in proj.get("bullets", []):
                _bullet(b)

    # ── Personal Projects ───────────────────────────────────────────────
    if data.get("personal_projects"):
        _section("Personal Projects")
        for proj in data["personal_projects"]:
            _row(proj["name"], proj.get("period", ""))
            for b in proj.get("bullets", []):
                _bullet(b)

    # ── Affiliations ────────────────────────────────────────────────────
    if data.get("affiliations"):
        _section("Professional Affiliations")
        for aff in data["affiliations"]:
            _bullet(aff)

    doc.save(str(path))


# ─── PDF document ─────────────────────────────────────────────────────────────

def _make_pdf_resume(data: dict, path: Path):
    from fpdf import FPDF, XPos, YPos

    profile = _profile()["personal"]
    portfolio = profile.get("portfolio", "")

    ACCENT = (26, 86, 219)
    MUTED = (85, 101, 122)

    class PDF(FPDF):
        def header(self): pass
        def footer(self): pass

    pdf = PDF()
    pdf.set_margins(18, 18, 18)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=18)
    W = pdf.w - 36

    def sec(title):
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(*ACCENT)
        pdf.cell(W, 5.5, title.upper(), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_draw_color(180, 200, 240)
        pdf.line(18, pdf.get_y(), pdf.w - 18, pdf.get_y())
        pdf.ln(2)
        pdf.set_text_color(0, 0, 0)

    def row(left, right, bold=True):
        dw = 44
        tw = W - dw
        y0 = pdf.get_y()
        pdf.set_font("Helvetica", "B" if bold else "", 9.5)
        pdf.multi_cell(tw, 4.8, _p(left), new_x=XPos.RIGHT, new_y=YPos.TOP)
        y1 = pdf.get_y()
        pdf.set_xy(pdf.l_margin + tw, y0)
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(*MUTED)
        pdf.cell(dw, 4.8, _p(right), align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)
        if pdf.get_y() < y1:
            pdf.set_y(y1)

    def body(text, size=9.5, italic=False, muted=False):
        sty = "I" if italic else ""
        pdf.set_font("Helvetica", sty, size)
        if muted:
            pdf.set_text_color(*MUTED)
        pdf.multi_cell(W, 4.5, _p(text))
        pdf.set_text_color(0, 0, 0)

    def bullet(text, size=9.5):
        pdf.set_font("Helvetica", "", size)
        y = pdf.get_y()
        pdf.set_x(pdf.l_margin + 3)
        pdf.cell(5, 4.5, "-")
        pdf.set_xy(pdf.l_margin + 8, y)
        pdf.multi_cell(W - 8, 4.5, _p(text))

    def skill_row(label, value):
        pdf.set_font("Helvetica", "B", 9.5)
        lw = pdf.get_string_width(f"{label}:    ")
        pdf.cell(lw, 4.8, f"{label}:    ", new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.set_font("Helvetica", "", 9.5)
        pdf.multi_cell(W - lw, 4.8, _p(value))

    # ── Name block ──────────────────────────────────────────────────────
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(W, 9, profile["name"].upper(), align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(*MUTED)
    pdf.cell(W, 6, _p(data.get("title_subtitle", "")), align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(0, 0, 0)

    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*MUTED)
    pdf.cell(W, 4.5,
             f"Mobile: {profile['phone']}  |  Email: {profile['email']}  |  Address: {profile['location']}",
             align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(W, 4.5,
             f"LinkedIn: {profile['linkedin']}  |  GitHub: {profile['github']}  |  Portfolio: {portfolio}",
             align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(3)

    # ── Career Objective ────────────────────────────────────────────────
    sec("Career Objective")
    body(data["career_objective"])

    # ── Technical Skills (no table — plain rows, ATS safe) ──────────────
    sec("Technical Skills")
    for cat, val in data["skills"].items():
        skill_row(cat, val)

    # ── Work Experience ─────────────────────────────────────────────────
    sec("Work Experience")
    for exp in data["experience"]:
        pdf.ln(1)
        row(f"{exp['title']}  |  {exp['company']}", exp["period"])
        body(f"Role Description: {exp['role_description']}", size=9, italic=True, muted=True)
        for b in exp.get("bullets", []):
            bullet(b)

    # ── Education ───────────────────────────────────────────────────────
    sec("Education")
    for edu in data["education"]:
        pdf.ln(1)
        row(f"{edu['degree']}  |  {edu['institution']}", edu["period"])
        for b in edu.get("bullets", []):
            bullet(b)

    # ── Academic Projects ───────────────────────────────────────────────
    if data.get("academic_projects"):
        sec("Academic Projects")
        for proj in data["academic_projects"]:
            pdf.ln(1)
            row(f"{proj['name']}  |  {proj.get('context', '')}", proj.get("period", ""))
            for b in proj.get("bullets", []):
                bullet(b)

    # ── Personal Projects ───────────────────────────────────────────────
    if data.get("personal_projects"):
        sec("Personal Projects")
        for proj in data["personal_projects"]:
            pdf.ln(1)
            row(proj["name"], proj.get("period", ""))
            for b in proj.get("bullets", []):
                bullet(b)

    # ── Affiliations ────────────────────────────────────────────────────
    if data.get("affiliations"):
        sec("Professional Affiliations")
        for aff in data["affiliations"]:
            bullet(aff)

    pdf.output(str(path))


# ─── Cover letter ─────────────────────────────────────────────────────────────

def _make_word_cover(data: dict, job: dict, path: Path):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    profile = _profile()["personal"]
    portfolio = profile.get("portfolio", "")

    def _sp(para, before=0, after=6):
        pPr = para._p.get_or_add_pPr()
        s = OxmlElement("w:spacing")
        s.set(qn("w:before"), str(before * 20))
        s.set(qn("w:after"), str(after * 20))
        pPr.append(s)

    p = doc.add_paragraph()
    _sp(p, after=0)
    r = p.add_run(profile["name"])
    r.bold = True; r.font.size = Pt(12)

    p = doc.add_paragraph()
    _sp(p, after=0)
    p.add_run(f"{profile['phone']}  |  {profile['email']}  |  {profile['location']}").font.size = Pt(10)

    p = doc.add_paragraph()
    _sp(p, after=14)
    r = p.add_run(f"{profile['linkedin']}  |  {portfolio}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    p = doc.add_paragraph()
    _sp(p, after=8)
    p.add_run(date.today().strftime("%d %B %Y")).font.size = Pt(10)

    p = doc.add_paragraph()
    _sp(p, after=0)
    p.add_run("Hiring Manager").font.size = Pt(10)

    p = doc.add_paragraph()
    _sp(p, after=14)
    p.add_run(job.get("company") or "").font.size = Pt(10)

    p = doc.add_paragraph()
    _sp(p, after=14)
    r = p.add_run(f"Re: {data['subject']}")
    r.bold = True; r.font.size = Pt(10)

    for para_text in data["paragraphs"]:
        p = doc.add_paragraph()
        _sp(p, after=8)
        p.add_run(para_text).font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph()
    _sp(p, before=14, after=0)
    p.add_run("Yours sincerely,").font.size = Pt(10)

    p = doc.add_paragraph()
    _sp(p, before=18, after=0)
    r = p.add_run(profile["name"])
    r.bold = True; r.font.size = Pt(10)

    doc.save(str(path))


def _make_pdf_cover(data: dict, job: dict, path: Path):
    from fpdf import FPDF, XPos, YPos

    profile = _profile()["personal"]
    portfolio = profile.get("portfolio", "")
    MUTED = (85, 101, 122)
    ACCENT = (26, 86, 219)

    class PDF(FPDF):
        def header(self): pass
        def footer(self): pass

    pdf = PDF()
    pdf.set_margins(22, 22, 22)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=22)
    W = pdf.w - 44

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(W, 7, profile["name"], new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*MUTED)
    pdf.cell(W, 5, f"{profile['phone']}  |  {profile['email']}  |  {profile['location']}",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(*ACCENT)
    pdf.cell(W, 5, f"{profile['linkedin']}  |  {portfolio}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 10)
    pdf.cell(W, 5, date.today().strftime("%d %B %Y"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)
    pdf.cell(W, 5, "Hiring Manager", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(W, 5, job.get("company") or "", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(8)

    pdf.set_font("Helvetica", "B", 10)
    pdf.multi_cell(W, 5, _p(f"Re: {data['subject']}"))
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 10)
    for para in data["paragraphs"]:
        pdf.multi_cell(W, 6, _p(para))
        pdf.ln(4)

    pdf.ln(8)
    pdf.cell(W, 5, "Yours sincerely,", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(10)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(W, 5, profile["name"], new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(str(path))


# ─── Plain text for modal display ─────────────────────────────────────────────

def _resume_to_text(data: dict) -> str:
    profile = _profile()["personal"]
    portfolio = profile.get("portfolio", "")
    L = [
        profile["name"].upper(),
        data.get("title_subtitle", ""),
        f"Mobile: {profile['phone']}  |  Email: {profile['email']}  |  Address: {profile['location']}",
        f"LinkedIn: {profile['linkedin']}  |  GitHub: {profile['github']}  |  Portfolio: {portfolio}",
        "",
        "CAREER OBJECTIVE",
        "-" * 40,
        data["career_objective"],
        "",
        "TECHNICAL SKILLS",
        "-" * 40,
    ]
    for cat, val in data["skills"].items():
        L.append(f"{cat}:  {val}")
    L += ["", "RELEVANT EXPERIENCE", "-" * 40]
    for exp in data.get("relevant_experience", []):
        L += [
            f"\n{exp['title']}  |  {exp.get('company', '')}  [{exp['period']}]",
            f"  Role Description: {exp['role_description']}",
        ]
        for b in exp.get("bullets", []):
            L.append(f"  - {b}")
    if data.get("other_experience"):
        L += ["", "ADDITIONAL EXPERIENCE", "-" * 40]
        for exp in data["other_experience"]:
            L += [
                f"\n{exp['title']}  |  {exp.get('company', '')}  [{exp['period']}]",
                f"  Role Description: {exp['role_description']}",
            ]
            for b in exp.get("bullets", []):
                L.append(f"  - {b}")
    L += ["", "EDUCATION", "-" * 40]
    for edu in data["education"]:
        L.append(f"\n{edu['degree']}  |  {edu['institution']}  [{edu['period']}]")
        for b in edu.get("bullets", []):
            L.append(f"  - {b}")
    if data.get("academic_projects"):
        L += ["", "ACADEMIC PROJECTS", "-" * 40]
        for proj in data["academic_projects"]:
            L.append(f"\n{proj['name']}  |  {proj.get('context', '')}  [{proj.get('period', '')}]")
            for b in proj.get("bullets", []):
                L.append(f"  - {b}")
    if data.get("personal_projects"):
        L += ["", "PERSONAL PROJECTS", "-" * 40]
        for proj in data["personal_projects"]:
            L.append(f"\n{proj['name']}  [{proj.get('period', '')}]")
            for b in proj.get("bullets", []):
                L.append(f"  - {b}")
    if data.get("affiliations"):
        L += ["", "PROFESSIONAL AFFILIATIONS", "-" * 40]
        for aff in data["affiliations"]:
            L.append(f"  - {aff}")
    return "\n".join(L)


def _cover_to_text(data: dict) -> str:
    profile = _profile()["personal"]
    portfolio = profile.get("portfolio", "")
    L = [
        profile["name"],
        f"{profile['phone']}  |  {profile['email']}  |  {profile['location']}",
        f"{profile['linkedin']}  |  {portfolio}",
        "",
        date.today().strftime("%d %B %Y"),
        "",
        "Hiring Manager",
        "",
        f"Re: {data['subject']}",
        "",
    ]
    for para in data["paragraphs"]:
        L += [para, ""]
    L += ["Yours sincerely,", "", profile["name"]]
    return "\n".join(L)


# ─── Public API ───────────────────────────────────────────────────────────────

def generate_for_job(job_id: str, force: bool = False) -> dict:
    with _db() as conn:
        job = conn.execute("SELECT * FROM jobs WHERE job_id=?", [job_id]).fetchone()
    if not job:
        raise ValueError(f"Job {job_id} not found")
    job = dict(job)

    resume_docx = OUTPUT_DIR / f"{job_id}_resume.docx"
    if not force and job.get("resume_text") and resume_docx.exists():
        # Cached DOCX exists — ensure PDFs are present too
        pdf_warnings = []
        for stem in ("resume", "cover"):
            pdf_path  = OUTPUT_DIR / f"{job_id}_{stem}.pdf"
            docx_path = OUTPUT_DIR / f"{job_id}_{stem}.docx"
            if not pdf_path.exists() and docx_path.exists():
                try:
                    _to_pdf(docx_path, pdf_path)
                except Exception as e:
                    pdf_warnings.append(f"{stem}: {e}")
        return {
            "resume_text":       job["resume_text"],
            "cover_letter_text": job["cover_letter_text"],
            "has_files":         True,
            "pdf_warnings":      pdf_warnings,
        }

    profile = _profile()
    result = _call_claude(profile, job)

    resume_data = result["resume"]
    cover_data = result["cover_letter"]

    resume_text = _resume_to_text(resume_data)
    cover_text = _cover_to_text(cover_data)

    try:
        _fill_resume(resume_data, OUTPUT_DIR / f"{job_id}_resume.docx")
        _fill_cover(cover_data, job, OUTPUT_DIR / f"{job_id}_cover.docx")
    except PermissionError as e:
        fname = Path(str(e)).name if "\\" in str(e) else str(e)
        raise PermissionError(
            f"Close the file in Word/Excel then try again — {fname}"
        ) from None

    pdf_warnings = []
    for stem in ("resume", "cover"):
        try:
            _to_pdf(OUTPUT_DIR / f"{job_id}_{stem}.docx", OUTPUT_DIR / f"{job_id}_{stem}.pdf")
        except Exception as pdf_err:
            pdf_warnings.append(f"{stem}: {pdf_err}")

    with _db() as conn:
        conn.execute(
            "UPDATE jobs SET resume_text=?, cover_letter_text=? WHERE job_id=?",
            [resume_text, cover_text, job_id],
        )
        conn.commit()

    return {
        "resume_text": resume_text,
        "cover_letter_text": cover_text,
        "pdf_warnings": pdf_warnings,
        "has_files": True,
    }
